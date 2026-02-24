"""
File Management Service.
Handles file storage with project-based organization.
"""
import asyncio
import io
import mimetypes
import uuid
import shutil
from pathlib import Path
from typing import Optional, Tuple
from backend.config import settings
import logging
import aiofiles
import boto3

logger = logging.getLogger(__name__)


class FileService:
    """Service for managing uploaded files."""

    SUPPORTED_EXTENSIONS = ['.pdf', '.txt', '.docx']
    
    def __init__(self):
        """Initialize file service."""
        self.object_storage_provider = str(settings.object_storage_provider or "local").strip().lower()
        self.upload_dir = Path(settings.upload_dir)
        self.max_size_bytes = settings.max_file_size_mb * 1024 * 1024
        self._s3_bucket = str(settings.aws_s3_bucket or "").strip()
        self._s3_region = str(settings.aws_s3_region or "us-east-1").strip()
        self._s3_presign_expiry = max(60, int(settings.aws_s3_presign_expiry_seconds or 900))
        self._s3_client = self._build_s3_client() if self._is_s3_enabled() else None
        
        # Create upload directory if it doesn't exist
        if not self._is_s3_enabled():
            self.upload_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"File service initialized in local mode (upload_dir={self.upload_dir})")
        else:
            logger.info("File service initialized in S3 mode (bucket=%s, region=%s)", self._s3_bucket, self._s3_region)

    def _is_s3_enabled(self) -> bool:
        return self.object_storage_provider in {"aws_s3", "s3"}

    def _build_s3_client(self):
        if not self._s3_bucket:
            raise ValueError("AWS_S3_BUCKET must be set when OBJECT_STORAGE_PROVIDER=aws_s3")
        return boto3.client(
            "s3",
            region_name=self._s3_region,
            aws_access_key_id=str(settings.aws_access_key_id or "").strip() or None,
            aws_secret_access_key=str(settings.aws_secret_access_key or "").strip() or None,
            endpoint_url=str(settings.aws_s3_endpoint_url or "").strip() or None,
        )

    def generate_object_key(self, project_id: int, unique_filename: str) -> str:
        return f"projects/{project_id}/documents/{unique_filename}"

    @staticmethod
    def _as_s3_uri(bucket: str, key: str) -> str:
        return f"s3://{bucket}/{key}"

    @staticmethod
    def _parse_s3_uri(file_path: str) -> Tuple[str, str]:
        value = str(file_path or "").strip()
        if not value.startswith("s3://"):
            raise ValueError("Invalid S3 URI")
        body = value[5:]
        bucket, _, key = body.partition("/")
        if not bucket or not key:
            raise ValueError("Invalid S3 URI")
        return bucket, key
    
    def get_project_dir(self, project_id: int) -> Path:
        """
        Get directory path for project.
        
        Args:
            project_id: Project ID
            
        Returns:
            Path to project directory
        """
        project_dir = self.upload_dir / f"project_{project_id}"
        project_dir.mkdir(parents=True, exist_ok=True)
        return project_dir
    
    def generate_unique_filename(self, original_filename: str) -> str:
        """
        Generate unique filename while preserving extension.
        
        Args:
            original_filename: Original file name
            
        Returns:
            Unique filename
        """
        file_ext = Path(original_filename).suffix
        unique_id = uuid.uuid4().hex[:12]
        safe_name = Path(original_filename).stem[:50]  # Limit length
        return f"{safe_name}_{unique_id}{file_ext}"
    
    async def save_upload_file(
        self,
        file_content: bytes,
        filename: str,
        project_id: int
    ) -> tuple[str, str]:
        """
        Save uploaded file to project directory.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            project_id: Project ID
            
        Returns:
            Tuple of (unique_filename, file_path)
            
        Raises:
            ValueError: If file is too large
        """
        # Check file size
        if len(file_content) > self.max_size_bytes:
            raise ValueError(
                f"File too large ({len(file_content)} bytes). "
                f"Maximum size is {settings.max_file_size_mb}MB"
            )
        
        # Generate unique filename
        unique_filename = self.generate_unique_filename(filename)
        if self._is_s3_enabled():
            key = self.generate_object_key(project_id, unique_filename)
            content_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"

            def _put() -> None:
                assert self._s3_client is not None
                self._s3_client.put_object(
                    Bucket=self._s3_bucket,
                    Key=key,
                    Body=file_content,
                    ContentType=content_type,
                )

            await asyncio.to_thread(_put)
            uri = self._as_s3_uri(self._s3_bucket, key)
            logger.info("Saved file to S3: %s (%s bytes)", uri, len(file_content))
            return unique_filename, uri

        # Local filesystem fallback
        project_dir = self.get_project_dir(project_id)
        file_path = project_dir / unique_filename
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)

        logger.info(f"Saved file: {file_path} ({len(file_content)} bytes)")
        return unique_filename, str(file_path)

    async def generate_presigned_upload(
        self,
        *,
        project_id: int,
        filename: str,
        content_type: str,
    ) -> dict:
        if not self._is_s3_enabled():
            raise ValueError("Presigned upload is available only when OBJECT_STORAGE_PROVIDER=aws_s3")

        unique_filename = self.generate_unique_filename(filename)
        key = self.generate_object_key(project_id, unique_filename)
        normalized_content_type = str(content_type or "").strip() or (mimetypes.guess_type(filename)[0] or "application/octet-stream")

        def _presign() -> str:
            assert self._s3_client is not None
            return self._s3_client.generate_presigned_url(
                "put_object",
                Params={
                    "Bucket": self._s3_bucket,
                    "Key": key,
                    "ContentType": normalized_content_type,
                },
                ExpiresIn=self._s3_presign_expiry,
                HttpMethod="PUT",
            )

        upload_url = await asyncio.to_thread(_presign)
        return {
            "upload_url": upload_url,
            "file_key": key,
            "unique_filename": unique_filename,
            "content_type": normalized_content_type,
            "expires_in": self._s3_presign_expiry,
        }
    
    async def delete_file(self, file_path: str) -> bool:
        """
        Delete file from storage.
        
        Args:
            file_path: Path to file
            
        Returns:
            True if deleted successfully
        """
        try:
            if str(file_path or "").startswith("s3://"):
                bucket, key = self._parse_s3_uri(file_path)

                def _delete_s3() -> None:
                    assert self._s3_client is not None
                    self._s3_client.delete_object(Bucket=bucket, Key=key)

                await asyncio.to_thread(_delete_s3)
                logger.info("Deleted S3 object: %s", file_path)
                return True

            path = Path(file_path)
            if path.exists():
                path.unlink()
                logger.info(f"Deleted file: {file_path}")
                return True
            else:
                logger.warning(f"File not found: {file_path}")
                return False
        except Exception as e:
            logger.error(f"Error deleting file: {str(e)}")
            raise
    
    async def delete_project_files(self, project_id: int) -> bool:
        """
        Delete all files for a project.
        
        Args:
            project_id: Project ID
            
        Returns:
            True if deleted successfully
        """
        try:
            if self._is_s3_enabled():
                prefix = f"projects/{project_id}/documents/"

                def _delete_prefix() -> int:
                    assert self._s3_client is not None
                    paginator = self._s3_client.get_paginator("list_objects_v2")
                    deleted = 0
                    for page in paginator.paginate(Bucket=self._s3_bucket, Prefix=prefix):
                        contents = page.get("Contents") or []
                        if not contents:
                            continue
                        objects = [{"Key": obj["Key"]} for obj in contents if obj.get("Key")]
                        if objects:
                            self._s3_client.delete_objects(Bucket=self._s3_bucket, Delete={"Objects": objects})
                            deleted += len(objects)
                    return deleted

                deleted_count = await asyncio.to_thread(_delete_prefix)
                logger.info("Deleted %s S3 objects for project %s", deleted_count, project_id)
                return deleted_count > 0

            project_dir = self.get_project_dir(project_id)
            if project_dir.exists():
                shutil.rmtree(project_dir)
                logger.info(f"Deleted project directory: {project_dir}")
                return True
            return False
        except Exception as e:
            logger.error(f"Error deleting project files: {str(e)}")
            raise
    
    def validate_file(self, filename: str, file_size: int) -> tuple[bool, Optional[str]]:
        """
        Validate file before upload.
        
        Args:
            filename: File name
            file_size: File size in bytes
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file extension
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file type. Supported: {self.SUPPORTED_EXTENSIONS}"
        
        # Check file size
        if file_size > self.max_size_bytes:
            return False, f"File too large. Maximum size is {settings.max_file_size_mb}MB"
        
        return True, None

    @classmethod
    def get_supported_extensions(cls) -> list[str]:
        return list(cls.SUPPORTED_EXTENSIONS)

    async def extract_text(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if str(file_path or "").startswith("s3://"):
            bucket, key = self._parse_s3_uri(file_path)

            def _download() -> bytes:
                assert self._s3_client is not None
                result = self._s3_client.get_object(Bucket=bucket, Key=key)
                body = result.get("Body")
                return body.read() if body is not None else b""

            data = await asyncio.to_thread(_download)
            return await self._extract_text_from_bytes(ext=ext, data=data)

        if ext == '.pdf':
            return await self._load_pdf(file_path)
        if ext == '.txt':
            return await self._load_txt(file_path)
        if ext == '.docx':
            return await self._load_docx(file_path)
        raise ValueError(f"Unsupported file type: {ext}")

    async def _extract_text_from_bytes(self, ext: str, data: bytes) -> str:
        if ext == '.pdf':
            return await self._load_pdf_bytes(data)
        if ext == '.txt':
            return self._load_txt_bytes(data)
        if ext == '.docx':
            return await self._load_docx_bytes(data)
        raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    async def _load_pdf(file_path: str) -> str:
        from pypdf import PdfReader

        def extract() -> str:
            reader = PdfReader(file_path)
            parts: list[str] = []
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    parts.append(text)
            return "\n\n".join(parts)

        return await asyncio.to_thread(extract)

    @staticmethod
    async def _load_pdf_bytes(file_content: bytes) -> str:
        from pypdf import PdfReader

        def extract() -> str:
            reader = PdfReader(io.BytesIO(file_content))
            parts: list[str] = []
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    parts.append(text)
            return "\n\n".join(parts)

        return await asyncio.to_thread(extract)

    @staticmethod
    async def _load_txt(file_path: str) -> str:
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as handle:
                return await handle.read()
        except UnicodeDecodeError:
            async with aiofiles.open(file_path, 'r', encoding='latin-1') as handle:
                return await handle.read()

    @staticmethod
    def _load_txt_bytes(file_content: bytes) -> str:
        try:
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            return file_content.decode("latin-1", errors="replace")

    @staticmethod
    async def _load_docx(file_path: str) -> str:
        from docx import Document

        def extract() -> str:
            doc = Document(file_path)
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    parts.append(para.text)
            return "\n\n".join(parts)

        return await asyncio.to_thread(extract)

    @staticmethod
    async def _load_docx_bytes(file_content: bytes) -> str:
        from docx import Document

        def extract() -> str:
            doc = Document(io.BytesIO(file_content))
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    parts.append(para.text)
            return "\n\n".join(parts)

        return await asyncio.to_thread(extract)
