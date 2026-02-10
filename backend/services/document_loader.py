"""
Document Loader Service.
Handles loading and extracting text from various document formats.
"""
from typing import Optional
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentLoaderService:
    """Service for loading documents and extracting text."""
    
    @staticmethod
    async def load_document(file_path: str) -> str:
        """
        Load document and extract text content.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file type is not supported
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return await DocumentLoaderService._load_pdf(file_path)
        elif file_ext == '.txt':
            return await DocumentLoaderService._load_txt(file_path)
        elif file_ext == '.docx':
            return await DocumentLoaderService._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    @staticmethod
    async def _load_pdf(file_path: str) -> str:
        """
        Load PDF file and extract text.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            from pypdf import PdfReader
            import asyncio

            def extract_text() -> tuple[str, int]:
                reader = PdfReader(file_path)
                text_parts = []

                for page in reader.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text)

                return "\n\n".join(text_parts), len(reader.pages)

            full_text, page_count = await asyncio.to_thread(extract_text)
            logger.info(f"Extracted {len(full_text)} characters from PDF ({page_count} pages)")

            return full_text

        except Exception as e:
            logger.error(f"Error loading PDF: {str(e)}")
            raise
    
    @staticmethod
    async def _load_txt(file_path: str) -> str:
        """
        Load text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content
        """
        try:
            import aiofiles
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                text = await f.read()
            
            logger.info(f"Loaded {len(text)} characters from TXT file")
            return text
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                import aiofiles as _af
                async with _af.open(file_path, 'r', encoding='latin-1') as f:
                    text = await f.read()
                logger.info(f"Loaded {len(text)} characters from TXT file (latin-1 encoding)")
                return text
            except Exception as e:
                logger.error(f"Error loading TXT file: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Error loading TXT file: {str(e)}")
            raise
    
    @staticmethod
    async def _load_docx(file_path: str) -> str:
        """
        Load DOCX file and extract text.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            import asyncio

            def _extract_docx():
                doc = Document(file_path)
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                return "\n\n".join(text_parts), len(doc.paragraphs)

            full_text, para_count = await asyncio.to_thread(_extract_docx)
            logger.info(f"Extracted {len(full_text)} characters from DOCX ({para_count} paragraphs)")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error loading DOCX: {str(e)}")
            raise
    
    @staticmethod
    def get_supported_extensions() -> list:
        """Get list of supported file extensions."""
        return ['.pdf', '.txt', '.docx']
    
    @staticmethod
    def is_supported_file(filename: str) -> bool:
        """
        Check if file type is supported.
        
        Args:
            filename: File name or path
            
        Returns:
            True if supported
        """
        ext = Path(filename).suffix.lower()
        return ext in DocumentLoaderService.get_supported_extensions()
